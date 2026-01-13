from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .schemas import ExtractedRules


def _set_margins(doc: Document, rules: ExtractedRules) -> None:
    section = doc.sections[0]
    section.left_margin = Mm(rules.margin_left_mm)
    section.right_margin = Mm(rules.margin_right_mm)
    section.top_margin = Mm(rules.margin_top_mm)
    section.bottom_margin = Mm(rules.margin_bottom_mm)


def _set_default_font(doc: Document, rules: ExtractedRules) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = rules.font_name
    font.size = Pt(rules.font_size_pt)

    # фикс для Word: продублировать name в rFonts
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:ascii"), rules.font_name)
    rFonts.set(qn("w:hAnsi"), rules.font_name)
    rFonts.set(qn("w:eastAsia"), rules.font_name)
    rFonts.set(qn("w:cs"), rules.font_name)


def _set_line_spacing(doc: Document, rules: ExtractedRules) -> None:
    # для каждого нового параграфа будем задавать spacing вручную в шаблонных блоках
    # (студенческий MVP: достаточно)
    pass


def _add_page_number_footer(doc: Document, rules: ExtractedRules) -> None:
    if not rules.page_numbering:
        return

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.name = rules.font_name
    run.font.size = Pt(rules.page_number_font_size_pt)

    # поле PAGE
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def _add_heading(doc: Document, text: str, level: int, rules: ExtractedRules) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = rules.font_name
    run.font.size = Pt(14)

    # простая имитация уровней
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # межстрочный интервал
    p.paragraph_format.line_spacing = rules.line_spacing


def _add_body_placeholder(doc: Document, text: str, rules: ExtractedRules) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = rules.line_spacing
    p.paragraph_format.first_line_indent = Mm(12.5)  # абзацный отступ ~1.25 см
    for run in p.runs:
        run.font.name = rules.font_name
        run.font.size = Pt(rules.font_size_pt)

def _add_toc(doc: Document, rules: ExtractedRules) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = rules.line_spacing

    run = p.add_run()
    run.font.name = rules.font_name
    run.font.size = Pt(rules.font_size_pt)

    # поле оглавления (Word сам обновляет: ПКМ -> "Обновить поле")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run._r.append(fld)

def generate_report_template(rules: ExtractedRules, out_path: str) -> None:
    doc = Document()

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    _set_margins(doc, rules)
    _set_default_font(doc, rules)
    _add_page_number_footer(doc, rules)

    # Титульный лист (рыба)
    _add_heading(doc, "ТИТУЛЬНЫЙ ЛИСТ", 1, rules)
    _add_body_placeholder(doc, "[Автоматически/вручную заполнить: ВУЗ, факультет, кафедра, тема, ФИО, руководитель]", rules)
    doc.add_page_break()

    # Содержание
    _add_heading(doc, "СОДЕРЖАНИЕ", 1, rules)
    _add_body_placeholder(doc, "[Содержание будет сформировано средствами Word после заполнения]", rules)
    doc.add_page_break()

    # Введение
    _add_heading(doc, "ВВЕДЕНИЕ", 1, rules)
    _add_body_placeholder(doc, "[Актуальность, цель, задачи, объект, предмет, методы, практическая значимость]", rules)

    # Главы
    _add_heading(doc, "ГЛАВА 1. [Название главы 1]", 1, rules)
    _add_heading(doc, "1.1 [Название параграфа]", 2, rules)
    _add_body_placeholder(doc, "[Текст...]", rules)

    _add_heading(doc, "ГЛАВА 2. [Название главы 2]", 1, rules)
    _add_heading(doc, "2.1 [Название параграфа]", 2, rules)
    _add_body_placeholder(doc, "[Текст...]", rules)

    _add_heading(doc, "ГЛАВА 3. [Название главы 3]", 1, rules)
    _add_heading(doc, "3.1 [Название параграфа]", 2, rules)
    _add_body_placeholder(doc, "[Текст...]", rules)

    # Заключение
    _add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1, rules)
    _add_body_placeholder(doc, "[Выводы и результаты]", rules)

    # Источники
    _add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1, rules)
    _add_body_placeholder(doc, "[1] ...", rules)

    # Приложения
    _add_heading(doc, "ПРИЛОЖЕНИЯ", 1, rules)
    _add_body_placeholder(doc, "[Код, скриншоты, схемы и т.д.]", rules)

    doc.save(out_path)
